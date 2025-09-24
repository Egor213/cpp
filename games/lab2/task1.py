import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches

N = 15

low = 1 / (N + 1)
high = N + 1
round_order = 3

# Матрица выигрышей
matrix = np.random.uniform(low, high, (10, 10))
matrix = np.round(matrix, round_order)

# Вероятности для критерия Байеса
probabilities = np.random.rand(10)
probabilities /= probabilities.sum()
probabilities = np.round(probabilities, round_order)

# Коэффициент склонности к риску для Гурвица
alpha = round(1 / (N + 2), round_order)

# --- Критерий Байеса ---
bayes_values = matrix.dot(probabilities)
bayes_values = np.round(bayes_values, round_order)
best_bayes = np.argmax(bayes_values)

# --- Критерий Вальда ---
wald_values = np.min(matrix, axis=1)
wald_values = np.round(wald_values, round_order)
best_wald = np.argmax(wald_values)

# --- Критерий Максимакса ---
maximax_values = np.max(matrix, axis=1)
maximax_values = np.round(maximax_values, round_order)
best_maximax = np.argmax(maximax_values)

# --- Критерий Гурвица ---
row_max = np.max(matrix, axis=1)
row_min = np.min(matrix, axis=1)
hurwicz_values = alpha * row_max + (1 - alpha) * row_min
hurwicz_values = np.round(hurwicz_values, round_order)
best_hurwicz = np.argmax(hurwicz_values)

# --- Критерий Лапласса ---
laplace_values = np.mean(matrix, axis=1)
laplace_values = np.round(laplace_values, round_order)
best_laplace = np.argmax(laplace_values)

# --- Критерий Сэвиджа ---
col_max = np.max(matrix, axis=0)
risk_matrix = col_max - matrix
savage_values = np.max(risk_matrix, axis=1)
savage_values = np.round(savage_values, round_order)
best_savage = np.argmin(savage_values)

# --- Таблица результатов ---
results = pd.DataFrame({
    "Стратегия": [f"Стр {i}" for i in range(10)],
    "Байес": bayes_values,
    "Вальд": wald_values,
    "Максимакс": maximax_values,
    "Гурвиц": hurwicz_values,
    "Лапласс": laplace_values,
    "Сэвидж": savage_values
})

# --- Таблица лучших стратегий ---
summary = pd.DataFrame({
    "Критерий": ["Байес", "Вальд", "Максимакс", "Гурвиц", "Лапласс", "Сэвидж"],
    "Лучшая стратегия": [
        f"Стр {best_bayes}",
        f"Стр {best_wald}",
        f"Стр {best_maximax}",
        f"Стр {best_hurwicz}",
        f"Стр {best_laplace}",
        f"Стр {best_savage}"
    ],
    "Значение": [
        bayes_values[best_bayes],
        wald_values[best_wald],
        maximax_values[best_maximax],
        hurwicz_values[best_hurwicz],
        laplace_values[best_laplace],
        savage_values[best_savage]
    ]
})

# --- Создание docx ---
doc = Document()

doc.add_heading("Матрица выигрышей 10x10", level=1)
table = doc.add_table(rows=11, cols=10)
for i in range(10):
    for j in range(10):
        table.cell(i, j).text = str(matrix[i, j])

# Вероятности
doc.add_heading("Вероятности для критерия Байеса", level=1)
prob_table = doc.add_table(rows=1, cols=10)
for j in range(10):
    prob_table.cell(0, j).text = str(probabilities[j])

doc.add_paragraph(f"\nСклонность к риску (альфа) для критерия Гурвица: {alpha}")

# Результаты по критериям
doc.add_heading("Результаты по всем критериям", level=1)
table_results = doc.add_table(rows=1, cols=len(results.columns))
# Заголовки
for j, col in enumerate(results.columns):
    table_results.cell(0, j).text = col
# Данные
for i in range(len(results)):
    row_cells = table_results.add_row().cells
    for j, col in enumerate(results.columns):
        row_cells[j].text = str(results.iloc[i, j])

# Лучшие стратегии
doc.add_heading("Лучшие стратегии по критериям", level=1)
table_summary = doc.add_table(rows=1, cols=len(summary.columns))
for j, col in enumerate(summary.columns):
    table_summary.cell(0, j).text = col
for i in range(len(summary)):
    row_cells = table_summary.add_row().cells
    for j, col in enumerate(summary.columns):
        row_cells[j].text = str(summary.iloc[i, j])

doc.save("результаты.docx")
print("Файл 'результаты.docx' успешно сохранён!")
