from math import comb, log2
import pandas as pd
import matplotlib.pyplot as plt

# Параметры колоды
total_cards = 50   # всего карт - 1!!!
bad_cards = 18      # плохие карты
good_cards = total_cards - bad_cards
hand_size = 6      # размер руки игрока А

def compute_probabilities(num_bad_in_hand):
    N1 = comb(bad_cards-num_bad_in_hand, 6) + comb(bad_cards-num_bad_in_hand, 5) * comb(good_cards - (6 - num_bad_in_hand), 1) + comb(bad_cards-num_bad_in_hand, 4) * comb(good_cards - (6 - num_bad_in_hand), 2)
    N2 = comb(bad_cards-num_bad_in_hand, 3) * comb(good_cards - (6 - num_bad_in_hand), 3)
    N3 = comb(good_cards - (6 - num_bad_in_hand), 6) + comb(good_cards - (6 - num_bad_in_hand), 5) * comb(bad_cards-num_bad_in_hand, 1) + comb(good_cards - (6 - num_bad_in_hand), 4) * comb(bad_cards-num_bad_in_hand, 2)
    N0 = comb(total_cards - 6, 6)
    p1 = N1 / N0
    p2 = N2 / N0
    p3 = N3 / N0
    probs = [p1, p2, p3]
    H_AB = -sum(p * log2(p) for p in probs if p > 0)
    
    return p1, p2, p3, H_AB

results = []
for bad_in_hand in range(0, hand_size+1):
    p1, p2, p3, H_AB = compute_probabilities(bad_in_hand)
    results.append([bad_in_hand, round(p1, 4), round(p2, 4), round(p3, 4), round(H_AB, 4)])

df = pd.DataFrame(results, columns=['Плохие у A', 'p1', 'p2', 'p3', 'H_AB'])
print(df)

plt.figure(figsize=(8,5))
plt.plot(df['Плохие у A'], df['H_AB'], marker='o')
plt.xlabel('Количество плохих карт у игрока A')
plt.ylabel('Энтропия H_AB')
plt.title('Энтропия')
plt.grid(True)
plt.show()



from docx import Document
doc = Document()
doc.add_heading("Энтропия и вероятности для игрока A", level=1)

table = doc.add_table(rows=1, cols=len(df.columns))
table.style = "Table Grid"

hdr_cells = table.rows[0].cells
for i, col_name in enumerate(df.columns):
    hdr_cells[i].text = str(col_name)

for i in range(len(df)):
    row_cells = table.add_row().cells
    for j, col in enumerate(df.columns):
        row_cells[j].text = str(df.iloc[i, j])

doc.save("entropy_results.docx")
print("Файл 'entropy_results.docx' успешно сохранён!")